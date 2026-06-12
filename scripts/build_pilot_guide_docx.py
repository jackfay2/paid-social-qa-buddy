"""Render docs/qa_buddy_pilot_guide.md to a shareable .docx.

Requires python-docx (`pip install python-docx`). Re-run after editing the
markdown or filling the [TODO] blanks:

    python scripts/build_pilot_guide_docx.py

Handles the markdown this guide uses: H1/H2, paragraphs, bullet + numbered
lists, pipe tables, fenced code blocks, blockquotes, and inline **bold** /
`code` / *italic*. Any text containing "TODO" is highlighted yellow so the
fill-in blanks stand out.
"""
import os
import re

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
MD = os.path.join(REPO, "docs", "qa_buddy_pilot_guide.md")
OUT = os.path.join(REPO, "docs", "qa_buddy_pilot_guide.docx")

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+?`|\*[^*]+?\*)")


def shade(pr, fill):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def add_inline(p, text, mono=False, base_bold=False):
    for seg in INLINE.split(text):
        if not seg:
            continue
        bold, italic, m, t = base_bold, False, mono, seg
        if seg.startswith("**") and seg.endswith("**"):
            bold, t = True, seg[2:-2]
        elif seg.startswith("`") and seg.endswith("`"):
            t = seg[1:-1]  # render literals as normal prose, not monospace
        elif seg.startswith("*") and seg.endswith("*"):
            italic, t = True, seg[1:-1]
        run = p.add_run(t)
        run.bold = bold
        run.italic = italic
        if m:
            run.font.name = "Consolas"
        if "TODO" in t:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def is_block_start(s):
    return (
        s.startswith(("#", ">", "|"))
        or s.strip().startswith("```")
        or s.strip() == "---"
        or bool(re.match(r"^[-*] ", s))
        or bool(re.match(r"^\d+\. ", s))
    )


def cells(row):
    parts = re.split(r"(?<!\\)\|", row.strip())
    if parts and parts[0] == "":
        parts = parts[1:]
    if parts and parts[-1] == "":
        parts = parts[:-1]
    return [c.strip().replace("\\|", "|") for c in parts]


def main():
    lines = open(MD, encoding="utf-8").read().split("\n")
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for name, sz in (("Heading 1", 18), ("Heading 2", 14)):
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(sz)
        st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        if line.strip().startswith("```"):
            i += 1
            code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            # Render fenced blocks (the Slack template, the service-account email)
            # as a clean light callout box in the normal proportional font, NOT a
            # monospace terminal block.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            shade(p._p.get_or_add_pPr(), "F4F6F8")
            for j, cl in enumerate(code):
                if j > 0:
                    p.add_run().add_break()
                r = p.add_run(cl)
                if "TODO" in cl:
                    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
            continue
        if line.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            tbl = []
            while i < n and lines[i].startswith("|"):
                tbl.append(lines[i])
                i += 1
            header = cells(tbl[0])
            body = [cells(r) for r in tbl[2:]]
            t = doc.add_table(rows=1, cols=len(header))
            t.style = "Table Grid"
            for k, h in enumerate(header):
                c = t.rows[0].cells[k]
                add_inline(c.paragraphs[0], h, base_bold=True)
                shade(c._tc.get_or_add_tcPr(), "D5E8F0")
            for brow in body:
                rc = t.add_row().cells
                for k in range(len(header)):
                    add_inline(rc[k].paragraphs[0], brow[k] if k < len(brow) else "")
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.strip() == "---":
            i += 1
            continue
        if line.startswith(">"):
            quote = []
            while i < n and lines[i].startswith(">"):
                quote.append(lines[i].lstrip(">").strip())
                i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            shade(p._p.get_or_add_pPr(), "FFF8E1")
            add_inline(p, " ".join(quote))
            continue
        if re.match(r"^[-*] ", line):
            item = [line[2:].strip()]
            i += 1
            while i < n and lines[i].startswith(("  ", "\t")) and lines[i].strip():
                item.append(lines[i].strip())
                i += 1
            add_inline(doc.add_paragraph(style="List Bullet"), " ".join(item))
            continue
        if re.match(r"^\d+\. ", line):
            item = [re.sub(r"^\d+\. ", "", line).strip()]
            i += 1
            while i < n and lines[i].startswith(("  ", "\t")) and lines[i].strip():
                item.append(lines[i].strip())
                i += 1
            add_inline(doc.add_paragraph(style="List Number"), " ".join(item))
            continue
        if line.strip() == "":
            i += 1
            continue
        para = [line.strip()]
        i += 1
        while i < n and lines[i].strip() and not is_block_start(lines[i]):
            para.append(lines[i].strip())
            i += 1
        add_inline(doc.add_paragraph(), " ".join(para))

    doc.save(OUT)
    print("wrote", OUT)


if __name__ == "__main__":
    main()
